# utils/file_processors.py - FILE PROCESSING UTILITIES

import io
import re
from typing import List, Optional, Dict, Any
from pathlib import Path
import logging

# File processing imports
try:
    import PyPDF2
    from docx import Document
    import pandas as pd
    PDF_AVAILABLE = True
    DOCX_AVAILABLE = True
except ImportError as e:
    logging.warning(f"File processing libraries not available: {e}")
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class FileProcessor:
    """Advanced file processing for medical abstracts"""
    
    def __init__(self):
        self.supported_formats = ['.txt', '.pdf', '.docx', '.csv', '.xlsx']
        
    def process_file(self, file_content: bytes, filename: str) -> str:
        """
        Process uploaded file and extract text content
        
        Args:
            file_content: Raw file bytes
            filename: Original filename with extension
            
        Returns:
            Extracted text content
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.txt':
            return self._process_txt(file_content)
        elif file_ext == '.pdf':
            return self._process_pdf(file_content)
        elif file_ext == '.docx':
            return self._process_docx(file_content)
        elif file_ext in ['.csv', '.xlsx']:
            return self._process_spreadsheet(file_content, file_ext)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _process_txt(self, file_content: bytes) -> str:
        """Process plain text file"""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1', errors='ignore')
        except Exception as e:
            logger.error(f"Error processing TXT file: {e}")
            raise
    
    def _process_pdf(self, file_content: bytes) -> str:
        """Process PDF file and extract text"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
            
            # Clean and join text
            full_text = '\n'.join(text_content)
            return self._clean_extracted_text(full_text)
            
        except Exception as e:
            logger.error(f"Error processing PDF file: {e}")
            raise
    
    def _process_pdf_by_pages(self, file_content: bytes) -> List[str]:
        """Process PDF file and extract text from each page separately"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            page_texts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():  # Only include non-empty pages
                    cleaned_text = self._clean_extracted_text(page_text)
                    if len(cleaned_text) > 100:  # Minimum abstract length
                        page_texts.append(cleaned_text)
            
            return page_texts
            
        except Exception as e:
            logger.error(f"Error processing PDF by pages: {e}")
            raise
    
    def _process_docx(self, file_content: bytes) -> str:
        """Process DOCX file and extract text"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            # Join and clean text
            full_text = '\n'.join(text_content)
            return self._clean_extracted_text(full_text)
            
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            raise
    
    def _process_spreadsheet(self, file_content: bytes, file_ext: str) -> str:
        """Process CSV/Excel files containing abstracts"""
        try:
            if file_ext == '.csv':
                df = pd.read_csv(io.BytesIO(file_content))
            else:  # .xlsx
                df = pd.read_excel(io.BytesIO(file_content))
            
            # Look for common abstract columns
            abstract_columns = ['abstract', 'text', 'content', 'summary', 'description']
            
            text_content = []
            for col in df.columns:
                if any(keyword in col.lower() for keyword in abstract_columns):
                    # Found abstract column
                    for value in df[col].dropna():
                        if isinstance(value, str) and len(value) > 100:  # Minimum abstract length
                            text_content.append(value)
            
            if not text_content:
                # Fallback: concatenate all text columns
                for col in df.select_dtypes(include=['object']).columns:
                    for value in df[col].dropna():
                        if isinstance(value, str) and len(value) > 50:
                            text_content.append(f"{col}: {value}")
            
            return '\n\n---\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing spreadsheet file: {e}")
            raise
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers and headers/footers
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Clean up common PDF artifacts
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\%\&\@\#\$\+\=\<\>\~\`\|\\]', ' ', text)
        
        return text.strip()
    
    def batch_process_files(self, files: List[Dict[str, Any]]) -> List[str]:
        """
        Process multiple files in batch
        
        Args:
            files: List of file dictionaries with 'content' and 'name' keys
            
        Returns:
            List of extracted text content
        """
        results = []
        
        for file_info in files:
            try:
                content = self.process_file(file_info['content'], file_info['name'])
                results.append(content)
            except Exception as e:
                logger.error(f"Error processing file {file_info['name']}: {e}")
                # Add placeholder for failed files
                results.append(f"ERROR: Failed to process {file_info['name']}: {str(e)}")
        
        return results
    
    def validate_file(self, filename: str, file_size: int, max_size: int = 100 * 1024 * 1024) -> bool:
        """
        Validate file before processing
        
        Args:
            filename: Name of the file
            file_size: Size of file in bytes
            max_size: Maximum allowed file size (default 100MB)
            
        Returns:
            True if file is valid
        """
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Check file size
        if file_size > max_size:
            raise ValueError(f"File too large: {file_size} bytes (max: {max_size} bytes)")
        
        # Check filename
        if not filename or len(filename) > 255:
            raise ValueError("Invalid filename")
        
        return True


class AbstractExtractor:
    """Extract abstracts from various document formats"""
    
    def __init__(self):
        self.file_processor = FileProcessor()
    
    def extract_abstracts_from_text(self, text: str) -> List[str]:
        """
        Extract individual abstracts from a text containing multiple abstracts
        
        Args:
            text: Input text potentially containing multiple abstracts
            
        Returns:
            List of individual abstract texts
        """
        # Common abstract separators
        separators = [
            r'\n\s*Abstract\s*\d+\s*[:\-]?\s*\n',
            r'\n\s*\d+\.\s*[A-Z][^.]*\.\s*\n',
            r'\n\s*Title\s*:\s*',
            r'\n\s*---+\s*\n',
            r'\n\s*\*{3,}\s*\n'
        ]
        
        abstracts = [text]  # Start with full text
        
        for separator in separators:
            new_abstracts = []
            for abstract in abstracts:
                parts = re.split(separator, abstract, flags=re.IGNORECASE)
                new_abstracts.extend([part.strip() for part in parts if part.strip()])
            abstracts = new_abstracts
        
        # Filter out very short texts (likely not full abstracts)
        filtered_abstracts = []
        for abstract in abstracts:
            if len(abstract) > 200:  # Minimum abstract length
                filtered_abstracts.append(abstract)
        
        return filtered_abstracts if filtered_abstracts else [text]
    
    def extract_structured_data(self, text: str) -> Dict[str, Any]:
        """
        Extract structured metadata from abstract text
        
        Args:
            text: Abstract text
            
        Returns:
            Dictionary with extracted metadata
        """
        metadata = {}
        
        # Extract title
        title_match = re.search(r'(?:Title\s*:?\s*)(.*?)(?:\n|Author)', text, re.IGNORECASE | re.DOTALL)
        if title_match:
            metadata['title'] = title_match.group(1).strip()
        
        # Extract authors
        authors_match = re.search(r'(?:Authors?\s*:?\s*)(.*?)(?:\n|Journal|Abstract)', text, re.IGNORECASE | re.DOTALL)
        if authors_match:
            metadata['authors'] = authors_match.group(1).strip()
        
        # Extract journal
        journal_match = re.search(r'(?:Journal\s*:?\s*)(.*?)(?:\n|Publication|DOI)', text, re.IGNORECASE)
        if journal_match:
            metadata['journal'] = journal_match.group(1).strip()
        
        # Extract DOI
        doi_match = re.search(r'(?:DOI\s*:?\s*)(10\.\d+\/[^\s]+)', text, re.IGNORECASE)
        if doi_match:
            metadata['doi'] = doi_match.group(1).strip()
        
        # Extract study type
        study_type_match = re.search(r'(?:Study Type\s*:?\s*)(.*?)(?:\n|Abstract)', text, re.IGNORECASE)
        if study_type_match:
            metadata['study_type'] = study_type_match.group(1).strip()
        
        return metadata


# Export main classes
__all__ = ['FileProcessor', 'AbstractExtractor'] 